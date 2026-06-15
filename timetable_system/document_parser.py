"""
High-precision document parser for university timetable system.

Parses Excel, Word, and PowerPoint files to extract courses, lecturers, and rooms.
Returns JSON matching the specified schema with proper normalization and defaults.
"""

import re
import json
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

try:
    import openpyxl
    from openpyxl.utils import column_index_from_string
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False


class DocumentParser:
    """Main parser class for university timetable documents."""

    def __init__(self):
        self.all_lecturers = []
        # Header synonyms mapping
        self.code_headers = ['code', 'course code', 'courseid', 'cid', 'subject code', 'module code']
        self.name_headers = ['course name', 'name', 'title', 'subject', 'course', 'course title', 'subject name']
        self.lecturer_headers = ['lecturer', 'lecturer name', 'lecturer_name', 'staff', 'instructor', 'tutor', 'teacher', 'facilitator']
        self.group_headers = ['group', 'part', 'year', 'semester', 'section', 'class', 'part/level', 'level/part']
        self.level_headers = ['level', 'course level', 'course_level', 'year level']
        self.duration_headers = ['duration', 'hours', 'lecture hours', 'credit', 'credits', 'hrs', 'duration hours']
        self.students_headers = ['max students', 'students', 'enrollment', 'capacity', 'max', 'class size', 'number of students']
        self.department_headers = ['department', 'dept', 'faculty', 'school', 'division']
        self.room_headers = ['room', 'room code', 'venue', 'room name', 'location', 'classroom', 'lab', 'hall']

        # Helper patterns
        self.group_pattern = re.compile(r'^(\d+\.\d+)$')
        self.roman_numeral_pattern = re.compile(r'^\s*(I+|II+|III+|IV+|V+)\s*$', re.IGNORECASE)
        self.email_special_chars = {'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e', 'à': 'a', 'â': 'a', 'ä': 'a', 'ù': 'u', 'û': 'u', 'ü': 'u', 'î': 'i', 'ï': 'i', 'ô': 'o', 'ö': 'o', 'ç': 'c', 'ñ': 'n', 'ø': 'o', 'å': 'a', 'æ': 'ae'}

    def normalize_text(self, text: str) -> str:
        """Normalize text: trim whitespace, collapse multiple spaces, remove non-printables."""
        if not text:
            return ''
        text = text.strip()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\x00|\x0b|\x0c|\r\n?', ' ', text)
        return text

    def normalize_name(self, name: str) -> str:
        """Normalize name to Title Case."""
        if not name:
            return ''
        name = self.normalize_text(name)
        parts = name.split()
        normalized_parts = []
        for part in parts:
            if part.upper() in ['MR', 'MRS', 'MS', 'DR', 'PROF', 'REV', 'HON', 'FR', 'SR', 'PR', 'ST', 'LT', 'GEN', 'ADM', 'MIS', 'REP', 'GOV', 'PRE', 'PRES']:
                normalized_parts.append(part)
            else:
                normalized_parts.append(part.title())
        return ' '.join(normalized_parts)

    def extract_surname_from_name(self, name: str) -> str:
        """Extract surname from a name for email/username generation."""
        if not name:
            return ''

        name = self.normalize_text(name)
        name = re.sub(r'^(Mr|Mrs|Ms|Miss|Dr|Prof|Eng|Sir|Madam|Capt|Lt|Gen|Rep|Sgt|Pfc|Cpl|Rev|Hon|Fr|Sr|Pr|St|Lt|Adm|Mis|Rep|Gov|Pre|Pres)\s+', '', name, flags=re.IGNORECASE)

        parts = name.split()
        if not parts:
            return ''

        surname = parts[-1]

        if len(surname) == 1 and parts[-2:] and len(parts[-2]) == 1 and parts[-2].isalpha():
            surname = parts[-2] + surname

        return surname.lower()

    def generate_lecturer_email(self, full_name: str) -> Tuple[str, str]:
        """
        Generate lecturer email and username from full name.
        Returns (email, username)
        """
        if not full_name:
            return '', ''

        name = self.normalize_text(full_name)
        name = re.sub(r'^(Mr|Mrs|Ms|Miss|Dr|Prof|Eng|Sir|Madam|Capt|Lt|Gen|Rep|Sgt|Pfc|Cpl|Rev|Hon|Fr|Sr|Pr|St|Lt|Adm|Mis|Rep|Gov|Pre|Pres)\s+', '', name, flags=re.IGNORECASE)

        parts = name.split()
        if not parts:
            return '', ''

        surname = parts[-1]

        initial = ''
        if len(parts) > 1:
            first_part = parts[0].rstrip('.')
            if len(first_part) == 1 and first_part.isalpha():
                initial = first_part.lower()
            elif len(first_part) > 1:
                initial = first_part[0].lower()
            else:
                initial = surname[0].lower()
        else:
            initial = surname[0].lower()

        username = surname.lower()
        email_local = f"{initial}{surname.lower()}"
        email = f"{email_local}@buse.ac.zw"

        return email, username

    def parse_lecturer_name(self, name: str) -> Dict[str, str]:
        """Parse lecturer name to extract title, full name, etc."""
        if not name:
            return {'full_name': '', 'title': None, 'username': '', 'email': ''}

        original = self.normalize_text(name)
        title = None

        title_match = re.match(r'^(Mr|Mrs|Ms|Miss|Dr|Prof|Eng|Sir|Madam|Capt|Lt|Gen|Rep|Sgt|Pfc|Cpl|Rev|Hon|Fr|Sr|Pr|St|Lt|Adm|Mis|Rep|Gov|Pre|Pres)\s+(.*)$', original, re.IGNORECASE)
        if title_match:
            title = title_match.group(1).title()
            full_name = title_match.group(2)
        else:
            full_name = original

        email, username = self.generate_lecturer_email(full_name)

        return {
            'full_name': self.normalize_name(full_name),
            'title': title,
            'username': username,
            'email': email
        }

    def is_group_label(self, text: str) -> bool:
        """Check if text is a group label (e.g., 1.1, 1.2, 2.1, etc.)."""
        if not text:
            return False

        text = self.normalize_text(text)

        if self.group_pattern.match(text):
            return True

        if self.roman_numeral_pattern.match(text):
            return True

        return False

    def extract_group_from_context(self, row_values: List[str], col_index: int) -> Optional[str]:
        """Try to extract group label from surrounding cells in a row."""
        for i, val in enumerate(row_values):
            if i != col_index and val and self.is_group_label(val):
                return val
        return None

    def parse_excel_file(self, filepath: str) -> Dict[str, Any]:
        """Parse Excel file and return structured data."""
        if not OPENPYXL_AVAILABLE:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': 'openpyxl not installed'}

        self.all_lecturers = []  # Reset lecturers list
        result = {'courses': [], 'lecturers': [], 'rooms': [], 'source_file': filepath}

        try:
            wb = openpyxl.load_workbook(filepath, data_only=True)
        except Exception as e:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': f'Failed to load Excel file: {str(e)}'}

        for sheet_idx, ws in enumerate(wb.worksheets, 1):
            sheet_name = ws.title

            max_row = ws.max_row
            max_col = ws.max_column

            if max_row < 2 or max_col < 1:
                continue

            raw_rows = []
            for row_idx in range(1, min(max_row + 1, 1000)):
                row_values = []
                for col_idx in range(1, min(max_col + 1, 100)):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    value = cell.value
                    if value is None:
                        row_values.append('')
                    else:
                        row_values.append(str(value).strip())
                if any(v for v in row_values):
                    raw_rows.append(row_values)

            if not raw_rows:
                continue

            headers = raw_rows[0]
            data_rows = raw_rows[1:]

            col_map = self._map_columns(headers)

            for row_idx, row_values in enumerate(data_rows, 2):
                course_data = self._process_excel_row(row_values, col_map, sheet_name, row_idx)
                if course_data:
                    result['courses'].append(course_data)

            wb.close()

        # Add lecturers to result
        result['lecturers'] = self.all_lecturers
        return result

    def _map_columns(self, headers: List[str]) -> Dict[str, int]:
        """Map column headers to canonical field names."""
        col_map = {}
        for idx, header in enumerate(headers):
            header_norm = self.normalize_text(header).lower().replace('_', ' ').replace('-', ' ')

            if any(k in header_norm for k in self.code_headers):
                col_map['code'] = idx
            elif any(k in header_norm for k in self.name_headers):
                col_map['name'] = idx
            elif any(k in header_norm for k in self.lecturer_headers):
                col_map['lecturer'] = idx
            elif any(k in header_norm for k in self.group_headers):
                col_map['group'] = idx
            elif any(k in header_norm for k in self.level_headers):
                col_map['level'] = idx
            elif any(k in header_norm for k in self.duration_headers):
                col_map['duration'] = idx
            elif any(k in header_norm for k in self.students_headers):
                col_map['students'] = idx
            elif any(k in header_norm for k in self.department_headers):
                col_map['department'] = idx
            elif any(k in header_norm for k in self.room_headers):
                col_map['room'] = idx

        # Debug logging
        print(f"[PARSER DEBUG] Headers: {headers}")
        print(f"[PARSER DEBUG] Column map: {col_map}")
        return col_map

    def _process_excel_row(self, row_values: List[str], col_map: Dict[str, int], sheet_name: str, row_idx: int) -> Optional[Dict]:
        """Process a single Excel row and return course data if valid."""
        if not row_values:
            return None

        print(f"[PARSER DEBUG] Row {row_idx}: {row_values}")
        print(f"[PARSER DEBUG] Col map: {col_map}")

        code = row_values[col_map.get('code', -1)] if 'code' in col_map and col_map['code'] < len(row_values) else ''
        if not code:
            print(f"[PARSER DEBUG] No code found, skipping row")
            return None

        code = self.normalize_text(code)

        name = row_values[col_map.get('name', -1)] if 'name' in col_map and col_map['name'] < len(row_values) else ''
        if not name:
            name = code

        lecturer_text = row_values[col_map.get('lecturer', -1)] if 'lecturer' in col_map and col_map['lecturer'] < len(row_values) else ''
        group = row_values[col_map.get('group', -1)] if 'group' in col_map and col_map['group'] < len(row_values) else ''

        if not group:
            group = self.extract_group_from_context(row_values, col_map.get('group', -1))
            if not group:
                group = '1.1'

        duration_text = row_values[col_map.get('duration', -1)] if 'duration' in col_map and col_map['duration'] < len(row_values) else ''
        try:
            duration = int(float(duration_text)) if duration_text else 4
        except (ValueError, TypeError):
            duration = 4

        students_text = row_values[col_map.get('students', -1)] if 'students' in col_map and col_map['students'] < len(row_values) else ''
        try:
            students = int(students_text) if students_text else 50
        except (ValueError, TypeError):
            students = 50

        department = row_values[col_map.get('department', -1)] if 'department' in col_map and col_map['department'] < len(row_values) else ''

        level = row_values[col_map.get('level', -1)] if 'level' in col_map and col_map['level'] < len(row_values) else ''
        # Normalize level to valid values
        valid_levels = ['university', 'department', 'single']
        if level:
            level_lower = level.lower().strip()
            if level_lower in valid_levels:
                level = level_lower
            else:
                level = ''

        room = row_values[col_map.get('room', -1)] if 'room' in col_map and col_map['room'] < len(row_values) else ''
        if room:
            room_data = self._create_room_data(room, sheet_name, row_idx)
            if room_data:
                return room_data

        lecturer_username = ''
        lecturer_email = ''
        lecturer_note = ''

        if lecturer_text and lecturer_text.upper() not in ['TBA', 'TBC', 'N/A', '', '']:
            lecturer_name = self.normalize_text(lecturer_text)

            lecturer_info = self.parse_lecturer_name(lecturer_name)
            lecturer_username = lecturer_info['username']
            lecturer_email = lecturer_info['email']

            course_lecturer = {
                'full_name': lecturer_info['full_name'],
                'title': lecturer_info['title'],
                'username': lecturer_username,
                'email': lecturer_email,
                'password': '1234',
                'department': department,
                'notes': None,
                'source_location': f'sheet:{sheet_name};row:{row_idx}',
                'contact': None
            }

            existing_lecturer = next((l for l in self.all_lecturers if l['username'] == lecturer_username), None)
            if not existing_lecturer:
                self.all_lecturers.append(course_lecturer)
            else:
                lecturer_username = existing_lecturer['username']
                lecturer_email = existing_lecturer['email']
                lecturer_note = f"duplicate_lecturer:{lecturer_name}"
        else:
            if lecturer_text and lecturer_text.upper() in ['TBA', 'TBC', 'N/A', '']:
                lecturer_note = f"TBA ({lecturer_text})"

        notes = []
        if lecturer_note:
            notes.append(lecturer_note)

        course_data = {
            'course_code': code,
            'course_name': self.normalize_name(name) if name else code,
            'duration_hours': duration,
            'level': level,
            'department': self.normalize_name(department) if department else '',
            'color': 'blue',
            'max_students': students,
            'group': group,
            'lecturer_username': lecturer_username if lecturer_username else None,
            'lecturer_email': lecturer_email if lecturer_email else None,
            'notes': '; '.join(notes) if notes else None,
            'source_location': f'sheet:{sheet_name};row:{row_idx}',
            'flag': 'needs_lecturer' if lecturer_text and lecturer_text.upper() in ['TBA', 'TBC', 'N/A', ''] else None
        }

        return course_data

    def _create_room_data(self, room_code: str, sheet_name: str, row_idx: int) -> Optional[Dict]:
        """Create room data from room code."""
        room_code = self.normalize_text(room_code)
        if not room_code:
            return None

        capacity = None
        capacity_match = re.search(r'(\d+)', room_code)
        if capacity_match:
            capacity = int(capacity_match.group(1))

        return {
            'room_code': room_code,
            'capacity': capacity,
            'notes': None,
            'source_location': f'sheet:{sheet_name};row:{row_idx}'
        }

    def parse_word_file(self, filepath: str) -> Dict[str, Any]:
        """Parse Word document and return structured data."""
        if not DOCX_AVAILABLE:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': 'python-docx not installed'}

        result = {'courses': [], 'lecturers': [], 'rooms': [], 'source_file': filepath}
        self.all_lecturers = []

        try:
            doc = docx.Document(filepath)
        except Exception as e:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': f'Failed to load Word document: {str(e)}'}

        tables = doc.tables
        for table_idx, table in enumerate(tables, 1):
            rows = []
            for row in table.rows:
                row_values = []
                for cell in row.cells:
                    row_values.append(cell.text.strip())
                if any(v for v in row_values):
                    rows.append(row_values)

            if len(rows) < 2:
                continue

            headers = rows[0]
            data_rows = rows[1:]

            col_map = self._map_columns(headers)

            for row_idx, row_values in enumerate(data_rows, 2):
                course_data = self._process_table_row(row_values, col_map, f'slide:1;table:{table_idx}', row_idx)
                if course_data:
                    result['courses'].append(course_data)

        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        for para_idx, para in enumerate(paragraphs, 1):
            text = para.text.strip()
            if text and 'CS' in text and any(kw in text.lower() for kw in ['course', 'code', 'lecturer', 'duration']):
                course_data = self._process_text_paragraph(text, f'slide:1;page:1;para:{para_idx}')
                if course_data:
                    result['courses'].append(course_data)

        result['lecturers'] = self.all_lecturers
        return result

    def _process_table_row(self, row_values: List[str], col_map: Dict[str, int], source_loc: str, row_idx: int) -> Optional[Dict]:
        return self._base_process_row(row_values, col_map, source_loc, row_idx)

    def _process_text_paragraph(self, text: str, source_loc: str) -> Optional[Dict]:
        pattern = r'(\b\w+\d+\b)\s*-\s*(.+?)(?:\s*-\s*(\d+)\s*(?:hours|hrs|hours?))?(?:\s*-\s*(\d+)\s*students?)?(?:\s*-\s*(.+?))?$'
        match = re.search(pattern, text, re.IGNORECASE)

        if match:
            code = match.group(1)
            name = match.group(2) if match.group(2) else code
            duration = match.group(3) if match.group(3) else '4'
            max_students = match.group(4) if match.group(4) else '50'
            lecturer = match.group(5) if match.group(5) else ''

            col_map = {'code': 0, 'name': 1, 'lecturer': 4, 'duration': 2, 'students': 3}
            row_values = [code, name, '', '', duration, max_students, '', lecturer]

            return self._base_process_row(row_values, col_map, source_loc, 1)

        return None

    def _base_process_row(self, row_values: List[str], col_map: Dict[str, int], source_loc: str, row_idx: int) -> Optional[Dict]:
        if not row_values:
            return None

        code = row_values[col_map.get('code', -1)] if 'code' in col_map and col_map['code'] < len(row_values) else ''
        if not code:
            return None

        code = self.normalize_text(code)

        name = row_values[col_map.get('name', -1)] if 'name' in col_map and col_map['name'] < len(row_values) else ''
        if not name:
            name = code

        lecturer_text = row_values[col_map.get('lecturer', -1)] if 'lecturer' in col_map and col_map['lecturer'] < len(row_values) else ''

        group = col_map.get('group', -1) if 'group' in col_map and col_map['group'] < len(row_values) else ''
        if group:
            group = row_values[col_map['group']]
        else:
            group = '1.1'

        duration_text = row_values[col_map.get('duration', -1)] if 'duration' in col_map and col_map['duration'] < len(row_values) else ''
        try:
            duration = int(float(duration_text)) if duration_text else 4
        except (ValueError, TypeError):
            duration = 4

        students_text = row_values[col_map.get('students', -1)] if 'students' in col_map and col_map['students'] < len(row_values) else ''
        try:
            students = int(students_text) if students_text else 50
        except (ValueError, TypeError):
            students = 50

        level = row_values[col_map.get('level', -1)] if 'level' in col_map and col_map['level'] < len(row_values) else ''
        valid_levels = ['university', 'department', 'single']
        if level:
            level_lower = level.lower().strip()
            if level_lower in valid_levels:
                level = level_lower
            else:
                level = ''

        lecturer_username = ''
        lecturer_email = ''
        lecturer_note = ''

        if lecturer_text and lecturer_text.upper() not in ['TBA', 'TBC', 'N/A', '', '']:
            lecturer_name = self.normalize_text(lecturer_text)

            lecturer_info = self.parse_lecturer_name(lecturer_name)
            lecturer_username = lecturer_info['username']
            lecturer_email = lecturer_info['email']

            course_lecturer = {
                'full_name': lecturer_info['full_name'],
                'title': lecturer_info['title'],
                'username': lecturer_username,
                'email': lecturer_email,
                'password': '1234',
                'department': '',
                'notes': None,
                'source_location': source_loc,
                'contact': None
            }

            existing_lecturer = next((l for l in self.all_lecturers if l['username'] == lecturer_username), None)
            if not existing_lecturer:
                self.all_lecturers.append(course_lecturer)
            else:
                lecturer_username = existing_lecturer['username']
                lecturer_email = existing_lecturer['email']
                lecturer_note = f"duplicate_lecturer:{lecturer_name}"
        else:
            if lecturer_text and lecturer_text.upper() in ['TBA', 'TBC', 'N/A', '']:
                lecturer_note = f"TBA ({lecturer_text})"

        notes = []
        if lecturer_note:
            notes.append(lecturer_note)

        course_data = {
            'course_code': code,
            'course_name': self.normalize_name(name) if name else code,
            'duration_hours': duration,
            'level': level,
            'department': '',
            'color': 'blue',
            'max_students': students,
            'group': group,
            'lecturer_username': lecturer_username if lecturer_username else None,
            'lecturer_email': lecturer_email if lecturer_username else None,
            'notes': '; '.join(notes) if notes else None,
            'source_location': source_loc,
            'flag': 'needs_lecturer' if lecturer_text and lecturer_text.upper() in ['TBA', 'TBC', 'N/A', ''] else None
        }

        return course_data

    def parse_all_formats(self, filepath: str) -> Dict[str, Any]:
        """Parse document based on file extension."""
        ext = filepath.lower()
        if (ext.endswith('.xlsx') or ext.endswith('.xls')) and OPENPYXL_AVAILABLE:
            return self.parse_excel_file(filepath)

        if ext.endswith('.docx') and DOCX_AVAILABLE:
            return self.parse_word_file(filepath)

        if ext.endswith('.csv'):
            return self.parse_csv_file(filepath)

        return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': f'Unsupported or missing parser for {filepath}'}

    def parse_csv_file(self, filepath: str) -> Dict[str, Any]:
        """Parse CSV file and return structured data."""
        import csv
        import io
        
        self.all_lecturers = []
        result = {'courses': [], 'lecturers': [], 'rooms': [], 'source_file': filepath}

        try:
            with open(filepath, 'r', encoding='utf-8-sig') as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    content = f.read()
            except Exception as e:
                return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': f'Failed to read CSV file: {str(e)}'}
        except Exception as e:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': f'Failed to read CSV file: {str(e)}'}

        try:
            reader = csv.DictReader(io.StringIO(content))
            rows = list(reader)
        except Exception as e:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': f'Failed to parse CSV: {str(e)}'}

        if not rows:
            return {'courses': [], 'lecturers': [], 'rooms': [], 'notes': 'CSV file is empty'}

        # Map columns
        col_map = self._map_columns(list(rows[0].keys()))

        for row_idx, row in enumerate(rows, 2):
            # Convert DictReader row to list format expected by _process_excel_row
            row_values = [row.get(h, '') for h in reader.fieldnames]
            course_data = self._process_excel_row(row_values, col_map, 'CSV', row_idx)
            if course_data:
                result['courses'].append(course_data)

        result['lecturers'] = self.all_lecturers
        return result

    def export_to_json(self, parsed_data: Dict[str, Any]) -> str:
        """Export parsed data to JSON format matching the schema."""
        courses = []
        lecturers = []
        rooms = []

        for course in parsed_data.get('courses', []):
            course_record = {
                'course_code': course.get('course_code', ''),
                'course_name': course.get('course_name', ''),
                'duration_hours': course.get('duration_hours', 4),
                'level': course.get('level', ''),
                'department': course.get('department', ''),
                'color': course.get('color', 'blue'),
                'max_students': course.get('max_students', 50),
                'group': course.get('group') if course.get('group') else None,
                'lecturer_username': course.get('lecturer_username') if course.get('lecturer_username') else None,
                'lecturer_email': course.get('lecturer_email') if course.get('lecturer_email') else None,
                'notes': course.get('notes') if course.get('notes') else None,
                'source_location': course.get('source_location') if course.get('source_location') else None,
                'flag': course.get('flag') if course.get('flag') else None
            }
            courses.append(course_record)

        for lecturer in parsed_data.get('lecturers', []):
            lecturer_record = {
                'full_name': lecturer.get('full_name', ''),
                'title': lecturer.get('title') if lecturer.get('title') else None,
                'username': lecturer.get('username', ''),
                'email': lecturer.get('email', ''),
                'password': lecturer.get('password', '1234'),
                'department': lecturer.get('department') if lecturer.get('department') else None,
                'contact': lecturer.get('contact') if lecturer.get('contact') else None,
                'notes': lecturer.get('notes') if lecturer.get('notes') else None,
                'source_location': lecturer.get('source_location') if lecturer.get('source_location') else None
            }
            lecturers.append(lecturer_record)

        for room in parsed_data.get('rooms', []):
            room_record = {
                'room_code': room.get('room_code', ''),
                'capacity': room.get('capacity') if room.get('capacity') is not None else None,
                'notes': room.get('notes') if room.get('notes') else None,
                'source_location': room.get('source_location') if room.get('source_location') else None
            }
            rooms.append(room_record)

        return json.dumps({
            'courses': courses,
            'lecturers': lecturers,
            'rooms': rooms,
            'source_file': parsed_data.get('source_file', ''),
            'notes': parsed_data.get('notes', '')
        }, indent=2, ensure_ascii=False)


if __name__ == '__main__':
    parser = DocumentParser()
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_parser.py <filepath>")
        sys.exit(1)

    filepath = sys.argv[1]
    parsed_data = parser.parse_all_formats(filepath)
    json_output = parser.export_to_json(parsed_data)
    print(json_output)